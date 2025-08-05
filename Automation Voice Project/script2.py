# Converts docx files to audio files with grammar correction
# Requirements: gTTS, language_tool_python, langdetect, python-docx

import os
from pathlib import Path
from gtts import gTTS
import language_tool_python
from langdetect import detect
from docx import Document

# Paths
input_folder = Path("InputFiles")
output_text_folder = Path("OutputFiles")
output_audio_folder = Path("OutputAudio")

# Ensure output folders exist
output_text_folder.mkdir(exist_ok=True)
output_audio_folder.mkdir(exist_ok=True)

# Load grammar tool
tool = language_tool_python.LanguageTool('en-US')
lang_map = {
    "en": "en-US",      # English (US)
    "zh-cn": "zh-CN",   # Chinese (Simplified)
    "ms": "id",          # Malay → Indonesian, closest match
    }
docx_files = list(input_folder.glob("*.docx"))
# Loop through all .txt files in input folder
for file in docx_files:
    doc = Document(file)
    original_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])

    #Language Detection
    lang_code = detect(original_text)
    

    lt_lang = lang_map.get(lang_code)
    if lt_lang:
        tool = language_tool_python.LanguageTool(lt_lang)
        corrected_text = tool.correct(original_text)
    else:
        print(f"Grammar correction not available for {file.name}. Using original text.")
        corrected_text = original_text

    # Grammar check and correction
    corrected_text = tool.correct(original_text)

    # Save corrected text
    corrected_doc = Document()
    for paragraph in corrected_text.split("\n"):
        corrected_doc.add_paragraph(paragraph)

    corrected_filename = output_text_folder / f"{file.stem}_corrected.docx"
    corrected_doc.save(corrected_filename)
    print(f"Saved corrected DOCX to: {corrected_filename}")

    # Convert to voice
    tts = gTTS(corrected_text)
    audio_filename = output_audio_folder / f"{file.stem}_audio.mp3"
    tts.save(audio_filename)

    print(f"Processed: {file.name}")
