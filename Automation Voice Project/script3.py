import os
from pathlib import Path
from gtts import gTTS
from langdetect import detect
from docx import Document
import requests

# Paths
input_folder = Path("InputFiles")
output_text_folder = Path("OutputFiles")
output_audio_folder = Path("OutputAudio")

# Ensure output folders exist
output_text_folder.mkdir(exist_ok=True)
output_audio_folder.mkdir(exist_ok=True)

# LanguageTool API-supported codes
lang_map = {
    "en": "en-US",       # English
    "zh-cn": "zh-CN",    # Chinese (Simplified)
    "ms": "id",          # Malay → use Indonesian as fallback
}

def correct_grammar(text, lang_code):
    url = "https://api.languagetoolplus.com/v2/check"
    params = {
        "text": text,
        "language": lang_code
    }
    try:
        response = requests.post(url, data=params)
        result = response.json()
        corrected_text = text
        matches = sorted(result.get("matches", []), key=lambda m: m['offset'], reverse=True)

        for match in matches:
            offset = match['offset']
            length = match['length']
            replacement = match['replacements'][0]['value'] if match['replacements'] else None
            if replacement:
                corrected_text = corrected_text[:offset] + replacement + corrected_text[offset + length:]

        return corrected_text
    except Exception as e:
        print(f"Error using LanguageTool API: {e}")
        return text  # fallback: return original

# Process each .docx file
docx_files = list(input_folder.glob("*.docx"))

for file in docx_files:
    doc = Document(file)
    original_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
    print(f"\nProcessing: {file.name}")
    print(f"Original text preview: {original_text[:100]}...")

    # Detect language
    try:
        lang_code = detect(original_text)
        lt_lang = lang_map.get(lang_code, "en-US")  # default to English if unknown
        print(f"Detected language: {lang_code} Using LanguageTool language: {lt_lang}")
    except Exception as e:
        print(f"Language detection failed: {e}")
        lt_lang = "en-US"

    # Correct grammar using LanguageTool API
    corrected_text = correct_grammar(original_text, lt_lang)

    # Save corrected text
    corrected_doc = Document()
    for paragraph in corrected_text.split("\n"):
        corrected_doc.add_paragraph(paragraph)

    corrected_filename = output_text_folder / f"{file.stem}_corrected.docx"
    corrected_doc.save(corrected_filename)
    print(f"Saved corrected DOCX to: {corrected_filename}")

    # Convert to voice
    try:
        tts = gTTS(corrected_text, lang=lang_code)
        audio_filename = output_audio_folder / f"{file.stem}_audio.mp3"
        tts.save(audio_filename)
        print(f"Saved audio to: {audio_filename}")
    except Exception as e:
        print(f"TTS conversion failed: {e}")
